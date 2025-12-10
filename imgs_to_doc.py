#!/usr/bin/env python3
"""
Inserta todas las imágenes de una carpeta en un documento de Word (.docx),
una imagen por página.

Uso:
    python imagenes_a_docx.py /ruta/a/carpeta_de_imagenes salida.docx

Requisitos:
    pip install python-docx pillow
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Error: hace falta instalar python-docx. Ejecuta:")
    print("    pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    print("Error: hace falta instalar Pillow. Ejecuta:")
    print("    pip install pillow")
    sys.exit(1)

EXTENSIONES = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"}


def main():
    if len(sys.argv) != 3:
        print("Uso: python imagenes_a_docx.py <carpeta_imagenes> <salida.docx>")
        sys.exit(1)

    carpeta_imagenes = Path(sys.argv[1])
    salida_docx = Path(sys.argv[2])

    if not carpeta_imagenes.is_dir():
        print(f"Error: {carpeta_imagenes} no es una carpeta válida.")
        sys.exit(1)

    doc = Document()

    imagenes = [p for p in sorted(carpeta_imagenes.iterdir())
                if p.is_file() and p.suffix.lower() in EXTENSIONES]

    if not imagenes:
        print("No se encontraron imágenes en la carpeta.")
        sys.exit(1)

    for i, ruta in enumerate(imagenes):
        # Nueva página excepto la primera
        if i > 0:
            doc.add_page_break()

        doc.add_paragraph(ruta.name)

        # Redimensionar a un ancho máximo aproximado de 15 cm
        try:
            with Image.open(ruta) as img:
                width, height = img.size
        except Exception as e:
            print(f"Advertencia: no se pudo abrir {ruta}: {e}")
            continue

        # Ancho máximo en pulgadas (15 cm ~ 5.9 in)
        max_width_inches = 5.9

        # Insertar la imagen redimensionada
        par = doc.add_paragraph()
        run = par.add_run()
        run.add_picture(str(ruta), width=Inches(max_width_inches))

    salida_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida_docx)
    print(f"Documento Word generado en: {salida_docx}")


if __name__ == "__main__":
    main()
